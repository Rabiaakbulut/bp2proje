from docx import Document
#nesnesini oluşturup her türlü ekleme işlemini yapacağımız yapı
from docx.enum.text import WD_BREAK               #bos sayfa bırakma
from docx.shared import Pt                        #punto ayarı
from tkinter import*                              #arayüz işlemleri
from tkinter import messagebox                    #uyarı mesajı vermek
import os                                         #word dosyasını kaydetme
from docx.shared import Inches                    #şekil boyutu
from docx.enum.text import WD_ALIGN_PARAGRAPH     #paragraf hizalama
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT #şekil hizalama
from docx.enum.table import WD_TABLE_ALIGNMENT    #çizelge hizalama
import tkinter.font as font                       #buton font

document=Document()

girişroot=Tk()
girişroot.geometry("1800x850")
girişroot.title("TEZ YAZMA PROGRAMI")

butonfont=font.Font(family='Monaco',weight="bold",size=20)
butonfont1=font.Font(family='Monaco',weight="bold",size=5)

def yenipencere1():
    root = Tk()
    root.title("KAPAK - İTHAF - ÖNSÖZ SAYFALARI")
    root.geometry("1800x850")

    entry1 = Entry(root, width=20)
    entry1.grid(row=2, column=3)
    entry2 = Entry(root, width=20)
    entry2.grid(row=3, column=3)
    entry3 = Entry(root, width=20)
    entry3.grid(row=4, column=3)
    entry4 = Entry(root, width=20)
    entry4.grid(row=5, column=3)
    entry5 = Entry(root, width=20)
    entry5.grid(row=6, column=3)
    entry6 = Entry(root, width=20)
    entry6.grid(row=7, column=3)
    entry7 = Entry(root, width=20)
    entry7.grid(row=8, column=3)
    entry8 = Entry(root, width=20)
    entry8.grid(row=9, column=3)
    entry9 = Entry(root, width=20)
    entry9.grid(row=10, column=3)

    entry00 = Entry(root, width=20)
    entry00.grid(row=11, column=3)
    entry01 = Entry(root, width=20)
    entry01.grid(row=12, column=3)
    entry02 = Entry(root, width=20)
    entry02.grid(row=13, column=3)
    entry03 = Entry(root, width=20)
    entry03.grid(row=14, column=3)

    entry11 = Entry(root, width=20)
    entry11.grid(row=2, column=7)
    entry12 = Entry(root, width=20)
    entry12.grid(row=3, column=7)
    entry13 = Entry(root, width=20)
    entry13.grid(row=4, column=7)
    entry14 = Entry(root, width=20)
    entry14.grid(row=5, column=7)
    entry15 = Entry(root, width=20)
    entry15.grid(row=6, column=7)
    entry16 = Entry(root, width=20)
    entry16.grid(row=7, column=7)
    entry17 = Entry(root, width=20)
    entry17.grid(row=8, column=7)
    entry18 = Entry(root, width=20)
    entry18.grid(row=9, column=7)
    entry19 = Entry(root, width=20)
    entry19.grid(row=10, column=7)

    entry20 = Entry(root, width=20)
    entry20.grid(row=3, column=11)
    entry21 = Entry(root, width=20)
    entry21.grid(row=2, column=11)

    etiketx = Label(root, text="").grid(row=0, column=0)
    etiket0 = Label(root, text="--- TEZ BİLGİLERİ ---", padx=50, pady=10, bg="thistle2", fg="purple4",font=butonfont).grid(row=1,
                                                                                                            column=1)
    etiket1 = Label(root, text="ENSTİTÜ : ", pady=10).grid(row=2, column=1)
    etiket2 = Label(root, text="TEZ ADI : ", pady=10).grid(row=3, column=1)
    etiket3 = Label(root, text="TEZ TİPİ : ", pady=10).grid(row=4, column=1)
    etiket4 = Label(root, text="AD : ", pady=10).grid(row=5, column=1)
    etiket5 = Label(root, text="SOYAD : ", pady=10).grid(row=6, column=1)
    etiket6 = Label(root, text="ANABİLİM DALI : ", pady=10).grid(row=7, column=1)
    etiket7 = Label(root, text="PROGRAM : ", pady=10).grid(row=8, column=1)
    etiket8 = Label(root, text="TESLİM TARİHİ (AY YIL) : ", pady=10).grid(row=9, column=1)
    etiket9 = Label(root, text="ÖĞRENCİ NO : ", pady=10).grid(row=10, column=1)
    etiket00 = Label(root, text="SAVUNMA TARİHİ : ", pady=10).grid(row=11, column=1)
    etiket01 = Label(root, text="TESLİM TARİHİ : ", pady=10).grid(row=12, column=1)
    etiket02 = Label(root, text="İTHAF EDİLEN KİŞİ : ", pady=10).grid(row=13, column=1)
    etiket03 = Label(root, text="MESLEK : ", pady=10).grid(row=14, column=1)

    etiket10 = Label(root, text="--- DANIŞMANLAR VE JÜRİLER ---", padx=50, pady=10, bg="thistle2", fg="purple4",font=butonfont).grid(
        row=1, column=5)
    etiket11 = Label(root, text="DANIŞMAN ÜNVAN :", pady=10).grid(row=2, column=5)
    etiket12 = Label(root, text="DANIŞMAN AD :", pady=10).grid(row=3, column=5)
    etiket13 = Label(root, text="DANIŞMAN SOYAD :", pady=10).grid(row=4, column=5)
    etiket14 = Label(root, text="JÜRİ1 ÜNVAN :", pady=10).grid(row=5, column=5)
    etiket15 = Label(root, text="JÜRİ1 AD :", pady=10).grid(row=6, column=5)
    etiket16 = Label(root, text="JÜRİ1 SOYAD :", pady=10).grid(row=7, column=5)
    etiket17 = Label(root, text="JÜRİ2 ÜNVAN :", pady=10).grid(row=8, column=5)
    etiket18 = Label(root, text="JÜRİ2 AD :", pady=10).grid(row=9, column=5)
    etiket19 = Label(root, text="JÜRİ2 SOYAD :", pady=10).grid(row=10, column=5)

    etiket20 = Label(root, text="--- ÖNSÖZ ---", padx=50, pady=10, bg="thistle2", fg="purple4",font=butonfont).grid(row=1, column=9)
    etiket21 = Label(root, text="TARİH (AY YIL)", pady=10).grid(row=2, column=9)
    etiket22 = Label(root, text="ÖNSÖZ", pady=10).grid(row=3, column=9)

    etiketx = Label(root, text="", padx=10, pady=10).grid(row=0, column=0)
    etiketxx = Label(root, text="", padx=10, pady=10).grid(row=0, column=2)
    etiketxxx = Label(root, text="", padx=10, pady=10).grid(row=0, column=4)
    etiketxxxx = Label(root, text="", padx=10, pady=10).grid(row=0, column=6)
    etiketxxxxx = Label(root, text="", padx=10, pady=10).grid(row=0, column=8)
    etiketxxxxxx = Label(root, text="", padx=10, pady=10).grid(row=1, column=10)
    etiketxxxxxxx = Label(root, text="", padx=10, pady=10).grid(row=1, column=12)


    def ad():
        girdi1 = entry1.get()
        girdi1 = girdi1.upper()
        enstitü = '\n', 'İSTANBUL TEKNİK ÜNİVERSİTESİ ★', girdi1, ' ENSTİTÜSÜ'
        baslık = document.add_paragraph()
        baslık.add_run(enstitü).underline = True  # altı çizili
        baslık.style.font.name = 'Times New Roman'
        baslık.style.font.size = Pt(12)
        baslık.style.font.bold = True
        baslık.alignment = 1  # .alignment =1 sola  =0 orta  =2 sağa yasla

        bosluk = document.add_paragraph('\n\n\n\n\n')

        girdi2 = entry2.get()
        girdi2 = girdi2.upper()
        tez_adı = document.add_paragraph(girdi2)
        tez_adı.alignment = 1

        bosluk = document.add_paragraph('\n\n\n\n\n\n\n\n\n')

        girdi3 = entry3.get()
        girdi3 = girdi3.upper()
        tezsecimi = document.add_paragraph(girdi3)
        tezsecimi.alignment = 1

        ad = entry4.get()
        soyad = entry5.get()
        ad = ad.title()
        soyad = soyad.upper()
        ad_soyad = ad + " " + soyad
        adsoyad = document.add_paragraph(ad_soyad + "\n\n\n\n\n")
        adsoyad.alignment = 1

        girdi6 = entry6.get()
        girdi6 = girdi6.title()
        anabilimdalı = document.add_paragraph(girdi6 + " Anabilim Dalı")
        anabilimdalı.alignment = 1

        girdi7 = entry7.get()
        girdi7 = girdi7.title()
        program = document.add_paragraph(girdi7 + " Programı\n\n\n\n\n\n")
        program.alignment = 1

        girdi8 = entry8.get()
        girdi8 = girdi8.upper()
        tarih = document.add_paragraph(girdi8)
        tarih.alignment = 1

        ################################################################################################## SAYFA1
        bossayfa = document.add_paragraph()
        run = bossayfa.add_run()
        run.add_break(WD_BREAK.PAGE)
        ################################################################################################## SAYFA2

        ##sayfa1e öğenci no ve tez danışamnı eklenmiş hali
        enstitü = '\n', 'İSTANBUL TEKNİK ÜNİVERSİTESİ ★', girdi1, ' ENSTİTÜSÜ'
        baslık = document.add_paragraph()
        baslık.add_run(enstitü).underline = True
        baslık.alignment = 1
        bosluk = document.add_paragraph('\n\n\n\n\n')

        tez_adı = document.add_paragraph(girdi2)
        tez_adı.alignment = 1

        bosluk = document.add_paragraph('\n\n\n\n\n\n\n\n\n')

        tezsecimi = document.add_paragraph(girdi3)
        tezsecimi.alignment = 1

        adsoyad = document.add_paragraph(ad_soyad)
        adsoyad.alignment = 1

        girdiogrno = entry9.get()
        ogrno = document.add_paragraph(girdiogrno + '\n\n\n')
        ogrno.alignment = 1

        anabilimdalı = document.add_paragraph(girdi6 + " Anabilim Dalı")
        anabilimdalı.alignment = 1

        program = document.add_paragraph(girdi7 + " Programı\n\n")
        program.alignment = 1

        girdiunvan = entry11.get()
        girdiunvan = girdiunvan.title()
        girditezdanısmanadı = entry12.get()
        girditezdanısmanadı = girditezdanısmanadı.title()
        girditezdanısmansoyadı = entry13.get()
        girditezdanısmansoyadı = girditezdanısmansoyadı.upper()
        danısman = girdiunvan + ' ' + girditezdanısmanadı + ' ' + girditezdanısmansoyadı
        danısmanekle = document.add_paragraph('Tez danışmanı: ' + danısman + '\n\n')
        danısmanekle.alignment = 1

        tarih = document.add_paragraph(girdi8)
        tarih.alignment = 1

        ################################################################################################## SAYFA3
        bossayfa = document.add_paragraph()
        run = bossayfa.add_run()
        run.add_break(WD_BREAK.PAGE)
        ################################################################################################## SAYFA4

        juri1unvan = entry14.get()
        juri1unvan = juri1unvan.title()
        juri1ad = entry15.get()
        juri1ad = juri1ad.title()
        juri1soyad = entry16.get()
        juri1soyad = juri1soyad.upper()
        juri1 = juri1unvan + ' ' + juri1ad + ' ' + juri1soyad
        juri2unvan = entry17.get()
        juri2unvan = juri2unvan.title()
        juri2ad = entry18.get()
        juri2ad = juri2ad.title()
        juri2soyad = entry19.get()
        juri2soyad = juri2soyad.upper()
        juri2 = juri2unvan + ' ' + juri2ad + ' ' + juri2soyad

        ################################################################################################## SAYFA5

        bosluk = document.add_paragraph("\n\n\n")
        sayfa4x = (
                    "İTÜ, " + girdi1 + " Enstitüsü’nün " + girdiogrno + " numaralı Yüksek Lisans Öğrencisi " + ad_soyad + ", ilgili yönetmeliklerin belirlediği gerekli tüm şartları yerine getirdikten sonra hazırladığı " + girdi2 + " başlıklı tezini aşağıda imzaları olan jüri önünde başarı ile sunmuştur.")
        sayfa4 = document.add_paragraph()
        sayfa4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_MED
        sayfa4.add_run(sayfa4x).bold = False

        bosluk = document.add_paragraph("\n\n\n\n")
        sayfa4_2 = document.add_paragraph("Tez Danışmanı :\t" + danısman)
        uniad = ("İstanbul Teknik Üniversitesi                 \t\t........................")
        uni = document.add_paragraph()
        uni.add_run(uniad).bold = False
        uni.alignment = 2

        bosluk = document.add_paragraph("\n\n\n")
        sayfa4_3 = document.add_paragraph("Jüri Üyeleri :\t\t" + juri1)
        uni2 = document.add_paragraph()
        uni2.add_run(uniad).bold = False
        uni2.alignment = 2

        sayfa4_4 = document.add_paragraph("\t\t\t" + juri2)
        uni3 = document.add_paragraph()
        uni3.add_run(uniad).bold = False
        uni3.alignment = 2

        teslimtarihi = entry00.get()
        teslimtarihi = teslimtarihi.title()
        savunmatarihi = entry01.get()
        savunmatarihi = savunmatarihi.title()

        sayfa4_5 = document.add_paragraph("\n\n\n\n\n\n\n\nTeslim Tarihi               :   " + teslimtarihi)
        sayfa4_5 = document.add_paragraph("Savunma Tarihi           :   " + savunmatarihi)

        ithaf = entry02.get()
        ithaf = ithaf.title()
        sayfa5 = document.add_paragraph(
            "\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n" + ithaf + "\n\n\n\n\n\n\n\n\n\n")
        sayfa5.alignment = 2

        önsöz = entry20.get()
        tarihönsöz = entry21.get()
        sayfa8 = document.add_paragraph("ÖNSÖZ")
        önsöz2 = "\n" + önsöz
        sayfa8_1 = document.add_paragraph()
        sayfa8_1.add_run(önsöz2).bold = False
        sayfa8_1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_MED
        sayfa8_2 = document.add_paragraph()
        sayfa8_2.add_run(tarihönsöz).bold = False
        sayfa8_2.alignment = 0
        sayfa8_2_2 = document.add_paragraph()
        sayfa8_2_2.add_run(ad_soyad).bold = False
        sayfa8_2_2.alignment = 2

        meslek = entry03.get()
        meslek = meslek.title()
        sayfa8_3 = document.add_paragraph()
        sayfa8_3.add_run(meslek).bold = False
        sayfa8_3.alignment = 2

        bossayfa = document.add_paragraph()
        run = bossayfa.add_run()
        run.add_break(WD_BREAK.PAGE)

        messagebox.showinfo(title="DURUM", message="Girdiğiniz bilgiler tez dosyanıza başarıyla eklendi!")

    butono = Button(root, text="Ekle", command=ad, padx=23, pady=20, bg="mistyrose3", fg="purple4",font=butonfont).grid(row=15,
                                                                                                          column=11)
etiketgirişboşluk=Label(girişroot,text="",padx=270,pady=50).grid(row=0,column=0)
pencere1buton=Button(girişroot,text="KAPAK - İTHAF - ÖNSÖZ",command=yenipencere1, padx=54, pady=20, bg="mistyrose3", fg="purple4",font=butonfont).grid(row=1,column=1)


def yenipencere2():
    root2 = Tk()
    root2.geometry("1000x300")
    root2.title("KISALTMALAR VE SEMBOLLER")

    pencere2boşluk = Label(root2, text="", padx=30, pady=40).grid(row=1, column=4)
    pencere2boşluk2 = Label(root2, text="", padx=30, pady=40).grid(row=2, column=4)

    etiketkısaltmaad = Label(root2, text="Kısaltma / Açılım", padx=50, pady=20,font=butonfont).grid(row=1, column=1)
    entrykısaltmaad = Entry(root2, width=20)
    entrykısaltmaad.grid(row=1, column=2)
    entrykısaltmatanım = Entry(root2, width=20)
    entrykısaltmatanım.grid(row=1, column=3)

    etiketsembolad = Label(root2, text="Sembol / Açılım", padx=50, pady=20,font=butonfont).grid(row=3, column=1)
    entrysembolad = Entry(root2, width=20)
    entrysembolad.grid(row=3, column=2)
    entrysemboltanım = Entry(root2, width=20)
    entrysemboltanım.grid(row=3, column=3)

    boşlukısaltma = document.add_paragraph("\n\n\n")
    boşlukısaltma.add_run("KISALTMALAR").bold=TRUE
    def kısaltmaekle():
        kısaltmaad = entrykısaltmaad.get()
        kısaltmatanım=entrykısaltmatanım.get()
        kısaltmatanımyeni="\t\t: "+kısaltmatanım
        kıslatmasatır=kısaltmaad+kısaltmatanımyeni
        kısaltmaadparagraf=document.add_paragraph()
        kısaltmaadparagraf.add_run(kısaltmaad).bold = TRUE
        kısaltmaadparagraf.add_run(kısaltmatanımyeni).bold = FALSE
        kısaltmaadparagraf.style.font.name = 'Times New Roman'
        kısaltmaadparagraf.style.font.size = Pt(12)
        entrykısaltmaad.delete(0, last=len(kısaltmaad) + 1)
        entrykısaltmatanım.delete(0, last=len(kısaltmatanım) + 1)
    butonkısaltma = Button(root2, text="Ekle", command=kısaltmaekle,padx=50, bg="mistyrose3", fg="purple4",font=butonfont).grid(row=1,column=5)

    def geçiş():
        bossayfa = document.add_paragraph()
        run = bossayfa.add_run()
        run.add_break(WD_BREAK.PAGE)
        boşluksembol = document.add_paragraph("\n\n\n")
        boşluksembol.add_run("SEMBOLLER").bold = TRUE
        messagebox.showinfo(title="DURUM", message="Semboller kısmına geçiş yaptınız!")
    butongeçiş=Button(root2,text="Evet",command=geçiş,padx=50, bg="mistyrose2", fg="purple4",font=butonfont).grid(row=2,column=5)
    etiketsembolad = Label(root2, text="Kısaltmalar kısmını bitirdiniz mi?",padx=100,font=butonfont).grid(row=2, column=1)
    def sembolekle():
        sembolad = entrysembolad.get()
        semboltanım=entrysemboltanım.get()
        semboltanımyeni="\t\t: "+semboltanım
        sembolsatır=sembolad+semboltanımyeni
        semboladparagraf = document.add_paragraph()
        semboladparagraf.add_run(sembolad).bold = TRUE
        semboladparagraf.add_run(semboltanımyeni).bold = FALSE
        semboladparagraf.style.font.name = 'Times New Roman'
        semboladparagraf.style.font.size = Pt(12)
        entrysembolad.delete(0, last=len(sembolad) + 1)
        entrysemboltanım.delete(0, last=len(semboltanım) + 1)
    butonsembol = Button(root2, text="Ekle", command=sembolekle,padx=50, bg="mistyrose3", fg="purple4",font=butonfont).grid(row=3,column=5)

pencere2buton=Button(girişroot,text="KISALTMALAR VE SEMBOLLER",command=yenipencere2, padx=30, pady=20, bg="mistyrose3", fg="purple4",font=butonfont).grid(row=2,column=1)


def yenipencere3():
    root3=Tk()
    root3.geometry("1200x400")
    root3.title("ÖZET SAYFALARI")

    özetboşluk = Label(root3, text="", padx=30, pady=20).grid(row=0, column=4)

    entryözetbaşlık = Entry(root3, width=100)
    entryözetbaşlık.grid(row=1, column=3)
    etiketözetbaşlık = Label(root3, text="Özet Başlığı", padx=50, pady=20,font=butonfont).grid(row=1, column=1)

    entryözet = Text(root3, width=75,height=10)
    entryözet.grid(row=2,column=3)
    etiketözet = Label(root3, text="Özet", padx=50, pady=50,font=butonfont).grid(row=2, column=1)

    def özetbaşlığıekle():

        bossayfa = document.add_paragraph()
        run = bossayfa.add_run()
        run.add_break(WD_BREAK.PAGE)

        özetyazısıboşluk = document.add_paragraph("\n\n\n")

        özetbaşlık = entryözetbaşlık.get()
        özetbaşlık = özetbaşlık.upper()
        özetbaşlıkparagraf = document.add_paragraph()
        özetbaşlıkparagraf.add_run(özetbaşlık).bold = TRUE
        özetbaşlıkparagraf.alignment = 1
        özetbaşlıkparagraf.style.font.name = 'Times New Roman'
        özetbaşlıkparagraf.style.font.size = Pt(12)
        entryözetbaşlık.delete(0, last=len(özetbaşlık) + 1)

        özetyazısı = document.add_paragraph()
        özetyazısı.add_run("ÖZET\n").bold = TRUE
        özetyazısı.alignment = 1

    butonözetbaşlık = Button(root3, text="Başlık Ekle", command=özetbaşlığıekle, padx=80, bg="mistyrose3", fg="purple4",font=butonfont).grid(row=1, column=5)



    def özetparagrafekle():
        özet = entryözet.get("1.0","end")
        özetparagraf = document.add_paragraph()
        özetparagraf.add_run(özet).bold = False
        özetparagraf.style.font.name = 'Times New Roman'
        özetparagraf.style.font.size = Pt(12)
        özetparagraf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_MED
        entryözet.delete('1.0', END)
    butonözet = Button(root3, text="Paragraf Ekle", command=özetparagrafekle, padx=70, bg="mistyrose3", fg="purple4",font=butonfont).grid(row=2,
                                                                                                                 column=5)
pencere3buton=Button(girişroot,text="ÖZET",command=yenipencere3, padx=190, pady=20, bg="mistyrose3", fg="purple4",font=butonfont).grid(row=3,column=1)



def yenipencere4():
    root1 = Tk()
    root1.geometry("1800x850")
    root1.title("GİRİŞ VE SONRAKİ SAYFALAR")

    etiketpencere4boşluk = Label(root1, text="", padx=90, pady=20).grid(row=0, column=0)

    entrybaşlık = Entry(root1, width=20)
    entrybaşlık.grid(row=1, column=2)
    etiketbaşlık = Label(root1, text="Başlık", padx=50, pady=20,font=butonfont).grid(row=1, column=1)

    entryaltbaşlık = Entry(root1, width=20)
    entryaltbaşlık.grid(row=2, column=2)
    etiketaltbaşlık = Label(root1, text="Altbaşlık", padx=50, pady=20,font=butonfont).grid(row=2, column=1)

    entryparagraf = Entry(root1, width=65)
    entryparagraf.grid(row=3, column=2,columnspan=3)
    etiketparagraf = Label(root1, text="Paragraf", padx=50, pady=20,font=butonfont).grid(row=3, column=1)

    entryşekil = Entry(root1, width=20)
    entryşekil.grid(row=4, column=2)
    entryşekilbaşlık = Entry(root1, width=20)
    entryşekilbaşlık.grid(row=4, column=3)
    entryşekiltanım = Entry(root1, width=20)
    entryşekiltanım.grid(row=4, column=4)
    etiketşekil = Label(root1, text="Şekil Ad/Şekil No/Tanım", padx=60, pady=20,font=butonfont).grid(row=4, column=1)

    entrysatır = Entry(root1, width=20)
    entrysatır.grid(row=5, column=2)
    entrysütun = Entry(root1, width=20)
    entrysütun.grid(row=5, column=3)
    entrydeğerler = Entry(root1, width=20)
    entrydeğerler.grid(row=5, column=4)
    etiketdeğerler = Label(root1, text="Satır Sayısı/Sütun Sayısı/Değerler", padx=60, pady=20,font=butonfont).grid(row=5, column=1)
    etiketdeğerler2 = Label(root1, text="Çizelge Ad/Tanım", padx=60,font=butonfont).grid(row=6, column=1)
    entryçizelgebaşlık = Entry(root1, width=20)
    entryçizelgebaşlık.grid(row=6, column=2)
    entryçizelgetanım = Entry(root1, width=20)
    entryçizelgetanım.grid(row=6, column=3)

    entrymaddeekle=Entry(root1,width=20)
    entrymaddeekle.grid(row=7,column=2)
    etiketmaddeekle= Label(root1, text="Madde (List Bullet)",padx=60, pady=40,font=butonfont).grid(row=7, column=1)

    row = Label(root1, text="", padx=30, pady=50).grid(row=0, column=5)
    column5 = Label(root1, text="", padx=30).grid(row=1, column=5)

    def başlıkekle():
        bossayfa = document.add_paragraph()
        run = bossayfa.add_run()
        run.add_break(WD_BREAK.PAGE)
        başlık = entrybaşlık.get()
        başlık = başlık.upper()
        başlık = "\n\n\n" + başlık + "\n"
        başlıkx = document.add_paragraph()
        başlıkx.add_run(başlık).bold = TRUE
        başlıkx.style.font.name = 'Times New Roman'
        başlıkx.style.font.size = Pt(12)
        entrybaşlık.delete(0, last=len(başlık) + 1)

    butonbaşlık = Button(root1, text="Ekle", command=başlıkekle, padx=80, bg="mistyrose3", fg="purple4",font=butonfont).grid(row=1,
                                                                                                             column=6)

    def altbaşlıkekle():
        altbaşlık = entryaltbaşlık.get()
        altbaşlık = altbaşlık.title()
        altbaşlıkx = document.add_paragraph()
        altbaşlıkx.style.font.name = 'Times New Roman'
        altbaşlıkx.style.font.size = Pt(12)
        altbaşlıkx.add_run(altbaşlık).bold = TRUE
        entryaltbaşlık.delete(0, last=len(altbaşlık) + 1)

    butonaltbaşlık = Button(root1, text="Ekle", command=altbaşlıkekle, padx=80, bg="mistyrose3", fg="purple4",font=butonfont).grid(
        row=2, column=6)

    def paragrafekle():
        paragraf = entryparagraf.get()
        paragrafx = document.add_paragraph()
        paragrafx.add_run(paragraf).bold = False
        paragrafx.style.font.name = 'Times New Roman'
        paragrafx.style.font.size = Pt(12)
        paragrafx.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_MED
        entryparagraf.delete(0, last=len(paragraf) + 1)
    butonparagraf = Button(root1, text="Ekle", command=paragrafekle, padx=80, bg="mistyrose3", fg="purple4",font=butonfont).grid(row=3,
                                                                                                                 column=6)

    def şekilekle():
        # şekil sayfaya ortalanmış bir şekilde ekleniyor
        şekil = entryşekil.get()
        şekilx = document.add_paragraph()
        şekilx.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        şekilrun = şekilx.add_run("")
        şekilrun.add_picture(şekil, width=Inches(5))
        entryşekil.delete(0, last=len(şekil) + 1)

        şekilbaşlık = entryşekilbaşlık.get()
        şekilbaşlık = şekilbaşlık.title()
        şekiltanım = entryşekiltanım.get()
        şekiltam = şekilbaşlık + şekiltanım
        şekiladlandırma = document.add_paragraph()
        şekiladlandırma.add_run(şekilbaşlık).bold = True
        şekiladlandırma.add_run(" : ").bold = True
        şekiladlandırma.add_run(şekiltanım).bold = False
        şekiladlandırma.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        entryşekilbaşlık.delete(0, last=len(şekilbaşlık) + 1)
        entryşekiltanım.delete(0, last=len(şekiltanım) + 1)

    butonşekiladlandırma = Button(root1, text="Ekle", command=şekilekle, padx=80, bg="mistyrose3", fg="purple4",font=butonfont).grid(
        row=4, column=6)

    def çizelgeekle():
        çizelgebaşlık = entryçizelgebaşlık.get()
        çizelgebaşlık = çizelgebaşlık.title()
        çizelgetanım = entryçizelgetanım.get()
        şekiltam = çizelgebaşlık + çizelgetanım
        çizelgeadlandırma = document.add_paragraph()
        çizelgeadlandırma.add_run(çizelgebaşlık).bold = True
        çizelgeadlandırma.add_run(" : ").bold = True
        çizelgeadlandırma.add_run(çizelgetanım).bold = False
        çizelgeadlandırma.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        çizelgeadlandırma.style.font.name = 'Times New Roman'
        çizelgeadlandırma.style.font.size = Pt(12)
        entryçizelgebaşlık.delete(0, last=len(çizelgebaşlık) + 1)
        entryçizelgetanım.delete(0, last=len(çizelgetanım) + 1)

        satırsayısı = int(entrysatır.get())
        sütunsayısı = int(entrysütun.get())
        değerlerget = entrydeğerler.get()
        değerler = değerlerget.split()
        table = document.add_table(rows=satırsayısı, cols=sütunsayısı)
        # table.alignment=WD_TABLE_ALIGNMENT.CENTER
        table.style = document.styles['Light Grid Accent 1']  #################################################
        m = 0
        for j in range(0, satırsayısı):
            hücreler = table.rows[j].cells
            for i in range(0, sütunsayısı):
                hücreler[i].text = değerler[m]
                m += 1
        entrysatır.delete(0, last=len(entrysatır.get()) + 1)
        entrysütun.delete(0, last=len(entrysütun.get()) + 1)
        entrydeğerler.delete(0, last=len(entrydeğerler.get()) + 1)

        çizelgesonuboşluk = document.add_paragraph()

    butontablo = Button(root1, text="Ekle", command=çizelgeekle, padx=80, bg="mistyrose3", fg="purple4",font=butonfont).grid(
        row=5, column=6)

    def maddeekle():
        madde=entrymaddeekle.get()
        maddeparagraf = document.add_paragraph()
        maddeparagraf.add_run(madde).bold = False
        maddeparagraf.style = 'List Bullet'
        maddeparagraf.style.font.name = 'Times New Roman'
        maddeparagraf.style.font.size = Pt(12)
        entrymaddeekle.delete(0, last=len(madde) + 1)

    butonmadde = Button(root1, text="Ekle", command=maddeekle, padx=80, bg="mistyrose3", fg="purple4",font=butonfont).grid(row=7, column=6)

    def hepsiniekle():
        başlıkekle()
        altbaşlıkekle()
        paragrafekle()
        şekilekle()
        çizelgeekle()
        maddeekle()
    butonhepsiniekle = Button(root1, text="Hepsini Ekle", command=hepsiniekle, padx=50, pady=20, bg="mistyrose2",font=butonfont,
                              fg="purple4").grid(row=8, column=6)
pencere4buton=Button(girişroot,text="GİRİŞ",command=yenipencere4, padx=182, pady=20, bg="mistyrose3", fg="purple4",font=butonfont).grid(row=4,column=1)

def yenipencere5():
    root4=Tk()
    root4.geometry("1000x300")
    root4.title("KAYNAKLAR")

    bossayfa = document.add_paragraph()
    run = bossayfa.add_run()
    run.add_break(WD_BREAK.PAGE)

    kaynakyazısıboşluk = document.add_paragraph()
    kaynakyazısıboşluk.add_run("\n\n\nKAYNAKLAR").bold = TRUE
    kaynakboşluk = Label(root4, text="", padx=30, pady=20).grid(row=0, column=4)

    entrykaynakkalın = Entry(root4, width=60)
    entrykaynakkalın .grid(row=1, column=3)
    etiketkaynakkalın = Label(root4, text="Kaynak Başlık", padx=50, pady=20,font=butonfont).grid(row=1, column=1)

    entrykaynak= Entry(root4, width=60)
    entrykaynak.grid(row=2,column=3)
    etiketkaynak = Label(root4, text="Kaynak", padx=50, pady=50,font=butonfont).grid(row=2, column=1)


    def kaynakekle():
        kalınkaynak = entrykaynakkalın.get()
        kalınkaynak = kalınkaynak.title()
        kaynak=entrykaynak.get()
        kaynak=" "+kaynak
        kaynakparagraf = document.add_paragraph()
        kaynakparagraf.add_run(kalınkaynak).bold = TRUE
        kaynakparagraf.add_run(kaynak).bold = False
        kaynakparagraf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_MED
        kaynakparagraf.style.font.name = 'Times New Roman'
        kaynakparagraf.style.font.size = Pt(12)
        entrykaynakkalın.delete(0, last=len(kalınkaynak) + 1)
        entrykaynak.delete(0, last=len(kaynak) + 1)
    butonkaynak = Button(root4, text="Ekle", command=kaynakekle, padx=80, bg="mistyrose3", fg="purple4",font=butonfont).grid(row=2, column=5)
pencere5buton=Button(girişroot,text="KAYNAKLAR",command=yenipencere5, padx=150, pady=20, bg="mistyrose3", fg="purple4",font=butonfont).grid(row=5,column=1)

def yenipencere6():

    root5 = Tk()
    root5.title("ÖZGEÇMİŞ")
    root5.geometry("1000x500")

    bossayfa = document.add_paragraph()
    run = bossayfa.add_run()
    run.add_break(WD_BREAK.PAGE)

    özgeçmiş = document.add_paragraph()
    özgeçmiş.add_run("\n\n\nÖZGEÇMİŞ\n").bold = TRUE
    özgeçmiş.style.font.name = 'Times New Roman'
    özgeçmiş.style.font.size = Pt(12)

    etiket0 = Label(root5, text="", pady=10,padx=20, font=butonfont).grid(row=0, column=4)
    etiket00 = Label(root5, text="", pady=10, padx=20, font=butonfont).grid(row=0, column=2)
    etiket000 = Label(root5, text="", pady=10, padx=20, font=butonfont).grid(row=0, column=0)

    etiket1 = Label(root5, text="AD SOYAD", pady=10,font=butonfont).grid(row=2, column=1)
    etiket2 = Label(root5, text="DOĞUM TARİHİ", pady=10,font=butonfont).grid(row=3, column=1)
    etiket3 = Label(root5, text="DOĞUM YERİ", pady=10,font=butonfont).grid(row=4, column=1)
    etiket4 = Label(root5, text="E-POSTA", pady=10,font=butonfont).grid(row=5, column=1)
    etiket5 = Label(root5, text="LİSANS", pady=10,font=butonfont).grid(row=6, column=1)
    etiket6 = Label(root5, text="FOTOĞRAF", pady=10,font=butonfont).grid(row=7, column=1)
    etiket7 = Label(root5, text="MESLEKİ DENEYİM VE ÖDÜLLER", pady=10,font=butonfont).grid(row=8, column=1)

    entry1 = Entry(root5, width=20)
    entry1.grid(row=2, column=3)
    entry2 = Entry(root5, width=20)
    entry2.grid(row=3, column=3)
    entry3 = Entry(root5, width=20)
    entry3.grid(row=4, column=3)
    entry4 = Entry(root5, width=20)
    entry4.grid(row=5, column=3)
    entry5 = Entry(root5, width=20)
    entry5.grid(row=6, column=3)
    entry6 = Entry(root5, width=20)
    entry6.grid(row=7, column=3)
    entry7 = Entry(root5, width=20)
    entry7.grid(row=8, column=3)

    def özgeçmişekle():
        isim=entry1.get()
        doğumtarihi = entry2.get()
        doğumyeri = entry3.get()
        eposta = entry4.get()
        lisans = entry5.get()
        resim = entry6.get()

        resimx = document.add_paragraph()
        resimx.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        resimrun = resimx.add_run("")
        resimrun.add_picture(resim, width=Inches(3))
        entry6.delete(0, last=len(resim) + 1)

        isim=isim.title()
        isimsatırı = document.add_paragraph()
        isimsatırı.add_run("\nAd-Soyad\t\t: ").bold = TRUE
        isimsatırı.add_run(isim).bold = FALSE
        entry1.delete(0, last=len(isim) + 1)

        dtsatırı = document.add_paragraph()
        dtsatırı.add_run("Doğum Tarihi\t: ").bold = TRUE
        dtsatırı.add_run(doğumtarihi).bold = FALSE
        entry2.delete(0, last=len(doğumtarihi) + 1)

        dysatırı = document.add_paragraph()
        dysatırı.add_run("Doğum Yeri\t\t: ").bold = TRUE
        dysatırı.add_run(doğumyeri).bold = FALSE
        entry3.delete(0, last=len(doğumyeri) + 1)

        epostasatır = document.add_paragraph()
        epostasatır.add_run("E-posta\t\t: ").bold = TRUE
        epostasatır.add_run(eposta).bold = FALSE
        entry4.delete(0, last=len(eposta) + 1)

        öğrenimdurumu = document.add_paragraph()
        öğrenimdurumu.add_run("\nÖĞRENİM DURUMU:").bold=True

        lisanssatırı = document.add_paragraph()
        lisanssatırı.add_run("Lisans\t\t: ").bold = True
        lisanssatırı.add_run(lisans).bold = False
        lisanssatırı.style = 'List Bullet'
        entry5.delete(0, last=len(lisans) + 1)

        meslekisatır = document.add_paragraph()
        meslekisatır.add_run("\nMESLEKİ DENEYİM VE ÖDÜLLER:").bold = True
    butonözgeçmiş = Button(root5, text="Hepsini Ekle", command=özgeçmişekle, padx=46, bg="mistyrose3", fg="purple4",
                             font=butonfont).grid(row=7, column=5)

    def deneyimlerlistesiekle():
        deneyim = entry7.get()
        deneyimler = document.add_paragraph()
        deneyimler.add_run(deneyim).bold = False
        deneyimler.style = 'List Bullet'
        entry7.delete(0, last=len(deneyim) + 1)
    butonözgeçmiş = Button(root5, text="Ekle", command=deneyimlerlistesiekle, padx=80, bg="mistyrose3", fg="purple4",
                               font=butonfont).grid(row=8, column=5)
pencere6buton=Button(girişroot,text="ÖZGEÇMİŞ",command=yenipencere6, padx=158, pady=20, bg="mistyrose3", fg="purple4",font=butonfont).grid(row=6,column=1)

girişroot.mainloop()
document.save('demo.docx')